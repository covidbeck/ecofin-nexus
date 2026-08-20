import json
import logging
import re
from io import BytesIO

from docx import Document
from PIL import Image
from PyPDF2 import PdfReader

from app.core.config import settings
from app.schemas.utility_bill import (
    BusinessType,
    RegionCode,
    UtilityBillExtractionSchema,
    UtilityBillInputSchema,
)

logger = logging.getLogger(__name__)

EXTRACTION_SYSTEM_PROMPT = (
    "Ты — AI-агент экстракции данных. Извлеки данные из квитанции за "
    "электроэнергию и верни СТРОГО валидный JSON без маркдауна и лишнего текста. "
    "Не считай, не округляй, не прогнозируй — только поля, явно указанные в документе. "
    "JSON строго такой формы: "
    '{"consumption_kwh": number, "total_cost_kzt": number, '
    '"business_type": "bakery"|"catering", "location": "astana"|"almaty", '
    '"billing_period_days": integer}. '
    "location: astana если Астана/Нур-Султан, иначе almaty. "
    "business_type: bakery для пекарни, иначе catering. "
    "billing_period_days — число дней расчётного периода (обычно 28–31)."
)

FALLBACK_DATA = UtilityBillInputSchema(
    total_kwh=3000.0,
    cost_kzt=45000.0,
    business_type=BusinessType.bakery,
    region=RegionCode.astana,
    days_in_month=30,
)

LOCATION_ALIASES = {
    "astana": RegionCode.astana,
    "астана": RegionCode.astana,
    "nur-sultan": RegionCode.astana,
    "nursultan": RegionCode.astana,
    "нур-султан": RegionCode.astana,
    "almaty": RegionCode.almaty,
    "алматы": RegionCode.almaty,
    "алмата": RegionCode.almaty,
}

BUSINESS_ALIASES = {
    "bakery": BusinessType.bakery,
    "пекарня": BusinessType.bakery,
    "пекарн": BusinessType.bakery,
    "catering": BusinessType.catering,
    "общепит": BusinessType.catering,
    "horeca": BusinessType.catering,
}

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_TYPE = "application/msword"
IMAGE_TYPES = {"image/png", "image/jpeg", "image/webp"}
FENCE_RE = re.compile(r"```(?:json)?\s*([\s\S]*?)\s*```", re.IGNORECASE)


class AIExtractorService:
    """Multimodal bill extraction via Gemini. Always returns a valid UtilityBillInputSchema."""

    async def extract_bill_data(self, payload: bytes, content_type: str) -> UtilityBillInputSchema:
        try:
            return self._extract_sync(payload, content_type)
        except Exception:
            logger.exception("Bill extraction failed; using fallback_data")
            return FALLBACK_DATA.model_copy()

    def _extract_sync(self, payload: bytes, content_type: str) -> UtilityBillInputSchema:
        if not payload:
            raise ValueError("empty upload")
        if not settings.gemini_api_key:
            raise ValueError("GEMINI_API_KEY is missing")

        try:
            import google.generativeai as genai
        except Exception as exc:
            raise RuntimeError("google-generativeai is unavailable") from exc

        genai.configure(api_key=settings.gemini_api_key)
        model_name = settings.gemini_model or "gemini-1.5-pro"
        model = genai.GenerativeModel(model_name, system_instruction=EXTRACTION_SYSTEM_PROMPT)
        generation_config = {"temperature": 0, "response_mime_type": "application/json"}
        contents = self._build_contents(payload, content_type)

        try:
            response = model.generate_content(contents, generation_config=generation_config)
            raw = (getattr(response, "text", None) or "").strip()
            payload_dict = self._parse_json(raw)
            extracted = UtilityBillExtractionSchema.model_validate(self._normalize(payload_dict))
            return extracted.to_input()
        except Exception:
            logger.exception("LLM call or JSON parse failed; using fallback_data")
            return FALLBACK_DATA.model_copy()

    def _build_contents(self, payload: bytes, content_type: str) -> list:
        if content_type == "application/pdf":
            bill_text = self._pdf_text(payload)
            if not bill_text.strip():
                raise ValueError("no extractable PDF text")
            return [f"Текст квитанции:\n{bill_text}"]
        if content_type in {DOCX_TYPE, DOC_TYPE}:
            bill_text = self._docx_text(payload)
            if not bill_text.strip():
                raise ValueError("no extractable DOCX text")
            return [f"Текст квитанции:\n{bill_text}"]
        if content_type in IMAGE_TYPES:
            image = Image.open(BytesIO(payload))
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGB")
            return ["Изображение квитанции за электроэнергию. Извлеки поля в JSON.", image]
        raise ValueError(f"unsupported content type: {content_type}")

    @staticmethod
    def _pdf_text(payload: bytes) -> str:
        reader = PdfReader(BytesIO(payload))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    @staticmethod
    def _docx_text(payload: bytes) -> str:
        document = Document(BytesIO(payload))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        return "\n".join(parts)

    @staticmethod
    def _parse_json(raw: str) -> dict:
        cleaned = (raw or "").strip()
        cleaned = cleaned.replace("```json", "").replace("```JSON", "").replace("```", "").strip()
        fenced = FENCE_RE.search(raw or "")
        if fenced:
            cleaned = fenced.group(1).strip()
        try:
            parsed = json.loads(cleaned)
        except Exception as exc:
            raise ValueError("LLM output is not valid JSON") from exc
        if not isinstance(parsed, dict):
            raise ValueError("LLM output is not a JSON object")
        return parsed

    @staticmethod
    def _normalize(payload: dict) -> dict:
        location = str(payload.get("location", "")).strip().lower()
        business = str(payload.get("business_type", "")).strip().lower()
        payload["location"] = LOCATION_ALIASES.get(location, location)
        payload["business_type"] = BUSINESS_ALIASES.get(business, business)
        return payload
