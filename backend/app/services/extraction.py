"""Bill extraction adapter (OCR/LLM). Extraction only — no math.

On any failure the adapter returns `needs_manual_entry=True`; it never
fabricates bill values. Every extracted field is a draft the user must review
and confirm before a consumption record is stored.
"""

import json
import logging
import re
from datetime import date
from io import BytesIO

from docx import Document
from PIL import Image
from PyPDF2 import PdfReader

from app.core.config import settings
from app.schemas.consumption import (
    BillUploadResponseSchema,
    ConsumptionCreateSchema,
    ExtractedFieldSchema,
)

logger = logging.getLogger(__name__)

EXTRACTION_SYSTEM_PROMPT = (
    "Ты — адаптер извлечения данных из счёта за электроэнергию. "
    "Верни СТРОГО валидный JSON без маркдауна и лишнего текста. "
    "Не считай, не округляй, не прогнозируй — только поля, явно указанные в документе. "
    "Если поле не найдено, ставь null. JSON строго такой формы: "
    '{"period_start": "YYYY-MM-DD"|null, "period_end": "YYYY-MM-DD"|null, '
    '"kwh": number|null, "total_cost_kzt": number|null, '
    '"fixed_charges_kzt": number|null}'
)

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_TYPE = "application/msword"
IMAGE_TYPES = {"image/png", "image/jpeg", "image/webp"}
FENCE_RE = re.compile(r"```(?:json)?\s*([\s\S]*?)\s*```", re.IGNORECASE)

MANUAL_ENTRY_MESSAGE = (
    "Не удалось надёжно извлечь данные из документа. Введите значения вручную — "
    "это безопаснее, чем работать с непроверенными цифрами."
)
REVIEW_MESSAGE = (
    "Черновик извлечения готов. Проверьте каждое поле и подтвердите данные — "
    "до подтверждения ничего не сохраняется."
)


class BillExtractionService:
    """Wraps the LLM call; every path ends in a reviewable draft or manual entry."""

    async def extract(self, payload: bytes, content_type: str) -> BillUploadResponseSchema:
        try:
            return self._extract_sync(payload, content_type)
        except Exception:
            logger.exception("Bill extraction failed; falling back to manual entry")
            return BillUploadResponseSchema(
                needs_manual_entry=True,
                fields=[],
                draft=None,
                warnings=["Извлечение недоступно (ошибка адаптера или нет API-ключа)."],
                message=MANUAL_ENTRY_MESSAGE,
            )

    def _extract_sync(self, payload: bytes, content_type: str) -> BillUploadResponseSchema:
        if not payload:
            raise ValueError("empty upload")
        if not settings.gemini_api_key:
            raise ValueError("GEMINI_API_KEY is missing")

        import google.generativeai as genai

        genai.configure(api_key=settings.gemini_api_key)
        model = genai.GenerativeModel(
            settings.gemini_model or "gemini-1.5-pro",
            system_instruction=EXTRACTION_SYSTEM_PROMPT,
        )
        contents = self._build_contents(payload, content_type)
        response = model.generate_content(
            contents,
            generation_config={"temperature": 0, "response_mime_type": "application/json"},
        )
        raw = (getattr(response, "text", None) or "").strip()
        parsed = self._parse_json(raw)
        return self._to_response(parsed)

    def _to_response(self, parsed: dict) -> BillUploadResponseSchema:
        fields: list[ExtractedFieldSchema] = []
        warnings: list[str] = []

        def field_of(name: str) -> object:
            value = parsed.get(name)
            fields.append(
                ExtractedFieldSchema(
                    name=name,
                    value=None if value is None else str(value),
                    confident=value is not None,
                )
            )
            return value

        period_start = field_of("period_start")
        period_end = field_of("period_end")
        kwh = field_of("kwh")
        cost = field_of("total_cost_kzt")
        fixed = field_of("fixed_charges_kzt")

        draft: ConsumptionCreateSchema | None = None
        try:
            if period_start and period_end and kwh is not None and cost is not None:
                draft = ConsumptionCreateSchema(
                    period_start=date.fromisoformat(str(period_start)),
                    period_end=date.fromisoformat(str(period_end)),
                    kwh=float(kwh),
                    cost_kzt=float(cost),
                    fixed_charges_kzt=float(fixed or 0.0),
                    data_quality="estimated",
                    source="upload",
                )
            else:
                warnings.append("Часть обязательных полей не найдена в документе.")
        except (ValueError, TypeError) as exc:
            warnings.append(f"Извлечённые значения не прошли валидацию: {exc}")
            draft = None

        return BillUploadResponseSchema(
            needs_manual_entry=draft is None,
            fields=fields,
            draft=draft,
            warnings=warnings,
            message=REVIEW_MESSAGE if draft else MANUAL_ENTRY_MESSAGE,
        )

    def _build_contents(self, payload: bytes, content_type: str) -> list:
        if content_type == "application/pdf":
            text = self._pdf_text(payload)
            if not text.strip():
                raise ValueError("no extractable PDF text")
            return [f"Текст счёта:\n{text}"]
        if content_type in {DOCX_TYPE, DOC_TYPE}:
            text = self._docx_text(payload)
            if not text.strip():
                raise ValueError("no extractable DOCX text")
            return [f"Текст счёта:\n{text}"]
        if content_type in IMAGE_TYPES:
            image = Image.open(BytesIO(payload))
            if image.mode not in {"RGB", "RGBA"}:
                image = image.convert("RGB")
            return ["Изображение счёта за электроэнергию. Извлеки поля в JSON.", image]
        raise ValueError(f"unsupported content type: {content_type}")

    @staticmethod
    def _pdf_text(payload: bytes) -> str:
        reader = PdfReader(BytesIO(payload))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    @staticmethod
    def _docx_text(payload: bytes) -> str:
        document = Document(BytesIO(payload))
        parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)

    @staticmethod
    def _parse_json(raw: str) -> dict:
        cleaned = (raw or "").strip()
        fenced = FENCE_RE.search(cleaned)
        if fenced:
            cleaned = fenced.group(1).strip()
        cleaned = cleaned.replace("```json", "").replace("```", "").strip()
        parsed = json.loads(cleaned)
        if not isinstance(parsed, dict):
            raise ValueError("LLM output is not a JSON object")
        return parsed
